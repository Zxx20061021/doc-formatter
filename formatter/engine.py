"""
公文格式识别规则引擎
基于 GB/T 9704-2012 标准
"""

import re
import copy
from docx import Document
from docx.shared import Pt, Mm, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import logging

logger = logging.getLogger(__name__)

# ──────────── 字体常量 ────────────
FONT_FANGSONG = "仿宋"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体"
FONT_XIAOBIAOSONG = "方正小标宋简体"

# ──────────── 字号常量 (pt) ────────────
SIZE_ERHAO = 22       # 二号 - 大标题
SIZE_SANHAO = 16      # 三号 - 正文/标题
SIZE_SIHAO = 14       # 四号 - 抄送

# ──────────── 行距常量 ────────────
LINE_SPACING = 28.8   # 磅 - 固定值

# ──────────── 页面格式 (mm) ────────────
PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)
MARGIN_TOP = Mm(37)
MARGIN_BOTTOM = Mm(35)
MARGIN_LEFT = Mm(28)
MARGIN_RIGHT = Mm(26)


class FormatEngine:
    """公文格式识别与自动排版引擎"""

    # ── 文种列表 ──
    WENZHONG = r"(通知|决定|意见|方案|办法|规定|报告|请示|批复|函|纪要|公告|通告|命令|令)"

    def __init__(self):
        self.confidence_marks = []  # 置信度标记

    def process(self, input_path, output_path):
        """主处理入口"""
        doc = Document(input_path)
        self.confidence_marks = []

        # 步骤0: 修复 LibreOffice 转换产生的 python-docx 不兼容对齐值
        self._fix_incompatible_alignments(doc)

        # 步骤1: 页面格式统一
        self._set_page_format(doc)

        # 步骤2: 识别大标题
        title_idx = self._identify_title(doc)

        # 步骤3: 识别发文字号
        doc_num_idx = self._identify_doc_number(doc, title_idx)

        # 步骤4: 识别主送机关
        recipient_idx = self._identify_recipient(doc, title_idx, doc_num_idx)

        # 步骤5: 逐段扫描正文
        body_start = self._get_body_start(title_idx, doc_num_idx, recipient_idx)
        self._scan_body(doc, body_start)

        # 步骤6: 识别落款区域
        self._identify_signature(doc)

        # 步骤7: 识别抄送
        self._identify_cc(doc)

        # 步骤8: 全局标点纠正
        self._fix_punctuation(doc)

        # 步骤9: 编号格式纠正
        self._fix_numbering(doc)

        doc.save(output_path)
        return {
            "confidence_marks": self.confidence_marks,
            "output_path": output_path
        }

    # ══════════════════════════════════════════════
    # 步骤1: 页面格式统一
    # ══════════════════════════════════════════════
    def _set_page_format(self, doc):
        """设置页面格式: A4, 上下左右边距, 行距"""
        for section in doc.sections:
            section.page_width = PAGE_WIDTH
            section.page_height = PAGE_HEIGHT
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT

    # ══════════════════════════════════════════════
    # 步骤2: 识别大标题
    # ══════════════════════════════════════════════
    def _identify_title(self, doc):
        """按优先级识别大标题"""
        best_idx = None
        best_priority = 999
        best_confidence = "low"

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 优先级1: 居中 + 字号最大
            if self._is_centered(para) and self._get_font_size(para) >= SIZE_ERHAO:
                if 1 < best_priority:
                    best_idx, best_priority, best_confidence = i, 1, "high"
                    continue

            # 优先级2: 居中 + 加粗
            if self._is_centered(para) and self._is_bold(para):
                if 2 < best_priority:
                    best_idx, best_priority, best_confidence = i, 2, "high"
                    continue

            # 优先级3: 居中 + 独立段落
            if self._is_centered(para) and self._is_isolated(doc, i):
                if 3 < best_priority:
                    best_idx, best_priority, best_confidence = i, 3, "medium"
                    continue

            # 优先级4: "关于"开头 + "的+文种"结尾
            if re.match(rf"^关于.+的{self.WENZHONG}$", text):
                if 4 < best_priority:
                    best_idx, best_priority, best_confidence = i, 4, "high"
                    continue

            # 优先级5: 文种结尾
            if re.match(rf".+{self.WENZHONG}$", text):
                if 5 < best_priority:
                    best_idx, best_priority, best_confidence = i, 5, "medium"
                    continue

            # 优先级6: 文档第一段 + 短文本
            if i == self._first_nonempty_para(doc) and len(text) <= 50 and not text.endswith("。"):
                if 6 < best_priority:
                    best_idx, best_priority, best_confidence = i, 6, "medium"
                    continue

            # 优先级7: 手动标记
            if text.startswith("#") or text.startswith("【"):
                clean = text.lstrip("#【】").strip()
                if clean:
                    if 7 < best_priority:
                        best_idx, best_priority, best_confidence = i, 7, "high"
                        # 去掉标记
                        self._set_para_text(para, clean)
                        continue

        if best_idx is not None:
            para = doc.paragraphs[best_idx]
            self._apply_title_format(para)
            self._add_confidence(best_idx, "大标题", best_confidence)

        return best_idx

    def _apply_title_format(self, para):
        """应用大标题格式: 方正小标宋简体, 二号, 居中, 行距28.8磅"""
        self._set_paragraph_font(para, FONT_XIAOBIAOSONG, SIZE_ERHAO, bold=False)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)

    # ══════════════════════════════════════════════
    # 步骤3: 识别发文字号
    # ══════════════════════════════════════════════
    def _identify_doc_number(self, doc, title_idx):
        """识别发文字号"""
        if title_idx is None:
            return None

        search_start = title_idx + 1
        # 只在标题后3段内查找
        for i in range(search_start, min(search_start + 3, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue

            # 匹配: 机关代字 + 年份 + 序号
            if re.search(r"[A-Za-z一-鿿]+[〔\[\(（][0-9]{4}[〕\]\)）][0-9]+号", text):
                para = doc.paragraphs[i]
                self._apply_doc_number_format(para)
                # 纠正年份括号
                corrected = self._fix_year_brackets(text)
                if corrected != text:
                    self._set_para_text(para, corrected)
                self._add_confidence(i, "发文字号", "high")
                return i

        return None

    def _apply_doc_number_format(self, para):
        """应用发文字号格式: 仿宋三号, 居中"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_line_spacing(para, LINE_SPACING)

    # ══════════════════════════════════════════════
    # 步骤4: 识别主送机关
    # ══════════════════════════════════════════════
    def _identify_recipient(self, doc, title_idx, doc_num_idx):
        """识别主送机关"""
        search_start = max(title_idx or 0, doc_num_idx or 0) + 1

        for i in range(search_start, min(search_start + 3, len(doc.paragraphs))):
            text = doc.paragraphs[i].text.strip()
            if not text:
                continue

            matched = False
            confidence = "low"

            # 优先级1: 标题后第一段 + 末尾冒号
            if text.endswith("：") or text.endswith(":"):
                matched = True
                confidence = "high"

            # 优先级2: 常见主送机关关键词开头
            if re.match(r"^(各省|各市|各县|各区|各部门|各单位|各分局|各直属|各位)", text):
                matched = True
                confidence = "high"

            # 优先级3: 末尾冒号的独立短段落
            if len(text) <= 30 and (text.endswith("：") or text.endswith(":")):
                matched = True
                confidence = "medium"

            if matched:
                para = doc.paragraphs[i]
                self._apply_recipient_format(para)
                # 确保末尾全角冒号
                if text.endswith(":"):
                    self._set_para_text(para, text[:-1] + "：")
                self._add_confidence(i, "主送机关", confidence)
                return i

        return None

    def _apply_recipient_format(self, para):
        """应用主送机关格式: 仿宋三号, 首行缩进2字符, 末尾全角冒号"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)

    # ══════════════════════════════════════════════
    # 步骤5: 逐段扫描正文
    # ══════════════════════════════════════════════
    def _get_body_start(self, title_idx, doc_num_idx, recipient_idx):
        """获取正文起始位置"""
        starts = [idx for idx in [title_idx, doc_num_idx, recipient_idx] if idx is not None]
        return max(starts) + 1 if starts else 0

    def _scan_body(self, doc, body_start):
        """逐段扫描正文，匹配标题/附件/正文格式"""
        # 先识别落款和抄送位置，以便排除
        signature_idx = self._find_signature_position(doc)
        cc_idx = self._find_cc_position(doc)

        for i in range(body_start, len(doc.paragraphs)):
            if i == signature_idx or i == cc_idx:
                continue
            # 跳过落款附近段落
            if signature_idx and i >= signature_idx - 1:
                if cc_idx and i < cc_idx:
                    continue
                elif not cc_idx and i >= signature_idx - 1 and i < len(doc.paragraphs) - 1:
                    continue

            text = doc.paragraphs[i].text.strip()
            if not text:
                continue

            para = doc.paragraphs[i]

            # 一级标题
            if self._match_heading1(para, text, i):
                self._apply_heading1_format(para)
                self._add_confidence(i, "一级标题", "high")
                continue

            # 二级标题
            if self._match_heading2(para, text, i):
                self._apply_heading2_format(para)
                self._add_confidence(i, "二级标题", "high")
                continue

            # 三级标题
            if self._match_heading3(para, text, i):
                self._apply_heading3_format(para)
                self._add_confidence(i, "三级标题", "high")
                continue

            # 四级标题
            if self._match_heading4(para, text, i):
                self._apply_heading4_format(para)
                self._add_confidence(i, "四级标题", "high")
                continue

            # 附件说明
            if self._match_attachment(para, text):
                self._apply_attachment_format(para)
                self._add_confidence(i, "附件说明", "high")
                continue

            # 默认: 正文格式
            self._apply_body_format(para)

    # ── 一级标题匹配 ──
    def _match_heading1(self, para, text, idx):
        # 优先级1: 中文数字 + 顿号
        if re.match(r"^[一二三四五六七八九十]+、", text):
            return True
        # 优先级2: "第" + 中文数字 + "部分/章/节"
        if re.match(r"^第[一二三四五六七八九十]+(部分|章|节)", text):
            return True
        # 优先级3: 手动标记
        if text.startswith("# "):
            self._set_para_text(para, text[2:].strip())
            return True
        if text.startswith("【一】"):
            return True
        # 优先级4: 加粗 + 独立短段落
        if self._is_bold(para) and len(text) <= 30 and self._is_isolated_standalone(para):
            return True
        return False

    def _apply_heading1_format(self, para):
        """一级标题: 黑体三号, 不加粗, 首行缩进2字符"""
        self._set_paragraph_font(para, FONT_HEITI, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)
        # 编号格式纠正: 确保用顿号
        text = para.text.strip()
        corrected = self._fix_heading1_numbering(text)
        if corrected != text:
            self._set_para_text(para, corrected)

    def _fix_heading1_numbering(self, text):
        """一级标题编号纠正"""
        # "一." → "一、"
        text = re.sub(r"^([一二三四五六七八九十]+)\.", r"\1、", text)
        # "1、" → "一、" (阿拉伯数字改中文数字)
        m = re.match(r"^([0-9]+)、", text)
        if m:
            num = self._arabic_to_chinese(int(m.group(1)))
            text = num + "、" + text[m.end():]
        # "第X " → "X、"
        m = re.match(r"^第([一二三四五六七八九十]+)\s", text)
        if m:
            text = m.group(1) + "、" + text[m.end():]
        # "(一）" → "一、" (括号编号改顿号)
        m = re.match(r"^[（(]([一二三四五六七八九十]+)[）)]", text)
        if m:
            text = m.group(1) + "、" + text[m.end():]
        return text

    # ── 二级标题匹配 ──
    def _match_heading2(self, para, text, idx):
        # 优先级1: 中文括号数字
        if re.match(r"^[（][一二三四五六七八九十]+[）]", text):
            return True
        # 优先级2: 半角括号中文数字 (需纠正)
        if re.match(r"^[\(][一二三四五六七八九十]+[\)]", text):
            return True
        # 优先级3: 手动标记
        if text.startswith("## "):
            self._set_para_text(para, text[3:].strip())
            return True
        if text.startswith("【二】"):
            return True
        # 优先级4: 加粗 + 独立短段落
        if self._is_bold(para) and len(text) <= 30:
            return True
        return False

    def _apply_heading2_format(self, para):
        """二级标题: 楷体三号, 不加粗, 首行缩进2字符"""
        self._set_paragraph_font(para, FONT_KAITI, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)
        # 编号纠正: 确保全角括号
        text = para.text.strip()
        corrected = self._fix_heading2_numbering(text)
        if corrected != text:
            self._set_para_text(para, corrected)

    def _fix_heading2_numbering(self, text):
        """二级标题编号纠正"""
        # 半角括号 → 全角括号
        text = re.sub(r"^\(([一二三四五六七八九十]+)\)", r"（\1）", text)
        # "1." → "（一）" (阿拉伯数字改中文括号)
        m = re.match(r"^([0-9]+)\.", text)
        if m:
            num = self._arabic_to_chinese(int(m.group(1)))
            text = f"（{num}）" + text[m.end():]
        # "一、" → "（一）" (顿号改全角括号，二级标题语境)
        m = re.match(r"^([一二三四五六七八九十]+)、", text)
        if m:
            text = f"（{m.group(1)}）" + text[m.end():]
        return text

    # ── 三级标题匹配 ──
    def _match_heading3(self, para, text, idx):
        # 优先级1: 阿拉伯数字 + 点
        if re.match(r"^[0-9]+\.", text):
            # 排除看起来像列表的（如 1. xxx 2. xxx 同一段）
            if not re.search(r"[0-9]+\.", text[2:]):
                return True
        # 优先级2: 阿拉伯数字 + 顿号 (需纠正)
        if re.match(r"^[0-9]+、", text):
            return True
        # 优先级3: 阿拉伯数字 + 逗号 (需纠正)
        if re.match(r"^[0-9]+，", text):
            return True
        # 优先级4: 阿拉伯数字 + 右括号 (需纠正)
        if re.match(r"^[0-9]+\)", text):
            return True
        # 优先级5: 手动标记
        if text.startswith("### "):
            self._set_para_text(para, text[4:].strip())
            return True
        if text.startswith("【三】"):
            return True
        return False

    def _apply_heading3_format(self, para):
        """三级标题: 仿宋三号, 加粗, 首行缩进2字符"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=True)
        self._set_first_line_indent(para, 2)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)
        # 编号纠正: 确保用点号
        text = para.text.strip()
        corrected = self._fix_heading3_numbering(text)
        if corrected != text:
            self._set_para_text(para, corrected)

    def _fix_heading3_numbering(self, text):
        """三级标题编号纠正"""
        # "1、" → "1."
        text = re.sub(r"^([0-9]+)、", r"\1.", text)
        # "1，" → "1."
        text = re.sub(r"^([0-9]+)，", r"\1.", text)
        # "1)" → "1."
        text = re.sub(r"^([0-9]+)\)", r"\1.", text)
        # "第一，" → "1."
        m = re.match(r"^第([一二三四五六七八九十]+)，", text)
        if m:
            num = self._chinese_to_arabic(m.group(1))
            text = f"{num}." + text[m.end():]
        return text

    # ── 四级标题匹配 ──
    def _match_heading4(self, para, text, idx):
        # 优先级1: 半角括号阿拉伯数字
        if re.match(r"^\([0-9]+\)", text):
            return True
        # 优先级2: 全角括号阿拉伯数字 (需纠正)
        if re.match(r"^（[0-9]+）", text):
            return True
        # 优先级3: 手动标记
        if text.startswith("#### "):
            self._set_para_text(para, text[5:].strip())
            return True
        if text.startswith("【四】"):
            return True
        return False

    def _apply_heading4_format(self, para):
        """四级标题: 仿宋三号, 不加粗, 首行缩进2字符"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)
        # 编号纠正: 全角括号→半角括号
        text = para.text.strip()
        corrected = self._fix_heading4_numbering(text)
        if corrected != text:
            self._set_para_text(para, corrected)

    def _fix_heading4_numbering(self, text):
        """四级标题编号纠正"""
        # 全角括号 → 半角括号
        text = re.sub(r"^（([0-9]+）)", r"(\1", text)
        text = re.sub(r"^（([0-9]+)）", r"(\1)", text)
        # "1." → "(1)" (点号改半角括号)
        m = re.match(r"^([0-9]+)\.", text)
        if m:
            text = f"({m.group(1)})" + text[m.end():]
        return text

    # ── 附件说明匹配 ──
    def _match_attachment(self, para, text):
        if re.match(r"^附件[：:]", text):
            return True
        return False

    def _apply_attachment_format(self, para):
        """附件说明: 仿宋三号, 左空2字"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)
        # 确保"附件"后用全角冒号
        text = para.text.strip()
        if text.startswith("附件:"):
            self._set_para_text(para, "附件：" + text[3:])

    # ══════════════════════════════════════════════
    # 步骤6: 识别落款
    # ══════════════════════════════════════════════
    def _find_signature_position(self, doc):
        """查找落款位置（日期所在段落）"""
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            text = doc.paragraphs[i].text.strip()
            if re.search(r"(20[0-9]{2})年([0-9]{1,2})月([0-9]{1,2})日", text):
                return i
        return None

    def _identify_signature(self, doc):
        """识别并格式化落款"""
        date_idx = self._find_signature_position(doc)
        if date_idx is None:
            return

        date_para = doc.paragraphs[date_idx]
        date_text = date_para.text.strip()

        # 格式化日期: 右对齐(右空2字), 阿拉伯数字, 月日不补零
        self._set_paragraph_font(date_para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_line_spacing(date_para, LINE_SPACING)

        # 纠正日期格式
        corrected = self._fix_date_format(date_text)
        if corrected != date_text:
            self._set_para_text(date_para, corrected)

        # 右空2字 (通过调整右缩进)
        self._set_right_indent(date_para, 2)

        self._add_confidence(date_idx, "成文日期", "high")

        # 日期上方: 发文机关署名
        if date_idx > 0:
            org_para = doc.paragraphs[date_idx - 1]
            org_text = org_para.text.strip()
            if org_text and len(org_text) <= 20 and not re.match(r"^[一二三四五六七八九十]+[、.]", org_text):
                self._set_paragraph_font(org_para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
                org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._set_line_spacing(org_para, LINE_SPACING)
                # 右空4字
                self._set_right_indent(org_para, 4)
                self._add_confidence(date_idx - 1, "发文机关署名", "medium")

    def _fix_date_format(self, text):
        """日期格式统一: 2025年7月16日（月日不补零）"""
        m = re.search(r"(20[0-9]{2})年(0?[0-9]{1,2})月(0?[0-9]{1,2})日", text)
        if m:
            year = m.group(1)
            month = str(int(m.group(2)))
            day = str(int(m.group(3)))
            new_date = f"{year}年{month}月{day}日"
            text = text[:m.start()] + new_date + text[m.end():]
        return text

    # ══════════════════════════════════════════════
    # 步骤7: 识别抄送
    # ══════════════════════════════════════════════
    def _find_cc_position(self, doc):
        """查找抄送位置"""
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            if re.match(r"^抄送[：:]", doc.paragraphs[i].text.strip()):
                return i
        return None

    def _identify_cc(self, doc):
        """识别并格式化抄送"""
        cc_idx = self._find_cc_position(doc)
        if cc_idx is None:
            return

        para = doc.paragraphs[cc_idx]
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SIHAO, bold=False)
        self._set_left_indent(para, 1)
        self._set_right_indent_chars(para, 1)
        self._set_line_spacing(para, LINE_SPACING)

        # 确保"抄送"后用全角冒号
        text = para.text.strip()
        if text.startswith("抄送:"):
            self._set_para_text(para, "抄送：" + text[3:])

        self._add_confidence(cc_idx, "抄送机关", "high")

    # ══════════════════════════════════════════════
    # 步骤8: 全局标点纠正
    # ══════════════════════════════════════════════
    def _fix_punctuation(self, doc):
        """全局标点符号纠正"""
        for para in doc.paragraphs:
            text = para.text
            if not text or not text.strip():
                continue

            changed = False
            new_text = text

            # 半角逗号 → 全角逗号（中文正文）
            if re.search(r"[一-鿿],", new_text):
                new_text = re.sub(r"(?<=[一-鿿]),", "，", new_text)
                changed = True

            # 半角句号 → 全角句号（中文正文，非编号）
            if re.search(r"[一-鿿]\.(?![0-9])", new_text):
                new_text = re.sub(r"(?<=[一-鿿])\.(?![0-9])", "。", new_text)
                changed = True

            # 半角冒号 → 全角冒号（中文正文）
            if re.search(r"[一-鿿]:", new_text):
                new_text = re.sub(r"(?<=[一-鿿]):", "：", new_text)
                changed = True

            # 半角分号 → 全角分号
            if re.search(r"[一-鿿];", new_text):
                new_text = re.sub(r"(?<=[一-鿿]);", "；", new_text)
                changed = True

            # 连续空格 → 单个空格
            if "  " in new_text:
                new_text = re.sub(r" {2,}", " ", new_text)
                changed = True

            # 多个回车 → ≤1个空行 (由段落间距处理)

            if changed:
                self._set_para_text(para, new_text)

        # 纠正发文字号年份括号
        for para in doc.paragraphs:
            text = para.text
            corrected = self._fix_year_brackets(text)
            if corrected != text:
                self._set_para_text(para, corrected)

    def _fix_year_brackets(self, text):
        """年份括号纠正: 方括号/圆括号 → 六角括号"""
        # [2025] → 〔2025〕
        text = re.sub(r"\[([0-9]{4})\]", r"〔\1〕", text)
        # （2025）→ 〔2025〕
        text = re.sub(r"（([0-9]{4})）", r"〔\1〕", text)
        # (2025) → 〔2025〕
        text = re.sub(r"\(([0-9]{4})\)", r"〔\1〕", text)
        return text

    # ══════════════════════════════════════════════
    # 步骤9: 编号格式纠正
    # ══════════════════════════════════════════════
    def _fix_numbering(self, doc):
        """编号格式纠正（已在各标题匹配中处理，此处做补充纠正）"""
        # 已在 apply_headingX_format 中处理
        pass

    # ══════════════════════════════════════════════
    # 正文格式
    # ══════════════════════════════════════════════
    def _apply_body_format(self, para):
        """应用正文格式: 仿宋三号, 两端对齐, 首行缩进2字符, 行距28.8磅"""
        self._set_paragraph_font(para, FONT_FANGSONG, SIZE_SANHAO, bold=False)
        self._set_first_line_indent(para, 2)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_line_spacing(para, LINE_SPACING)
        self._set_spacing_before_after(para, 0, 0)

    # ══════════════════════════════════════════════
    # 辅助方法
    # ══════════════════════════════════════════════
    def _fix_incompatible_alignments(self, doc):
        """修复 LibreOffice 转换产生的 python-docx 不兼容对齐值
        LibreOffice 使用 'start'/'end' 对齐，python-docx 不识别会报错
        """
        alignment_map = {
            'start': 'left',
            'end': 'right',
        }
        for para in doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    val = jc.get(qn('w:val'))
                    if val in alignment_map:
                        jc.set(qn('w:val'), alignment_map[val])

    def _is_centered(self, para):
        """判断段落是否居中"""
        try:
            return para.alignment == WD_ALIGN_PARAGRAPH.CENTER or \
                   (para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER)
        except (ValueError, KeyError):
            # LibreOffice 转换可能产生不兼容的对齐值，回退到检查 XML
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                jc = pPr.find(qn('w:jc'))
                if jc is not None:
                    return jc.get(qn('w:val')) in ('center', 'Center')
            return False

    def _is_bold(self, para):
        """判断段落是否加粗"""
        for run in para.runs:
            if run.bold:
                return True
        # 也检查 XML
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
            if rPr is not None:
                b = rPr.find(qn('w:b'))
                if b is not None:
                    return True
        return False

    def _get_font_size(self, para):
        """获取段落最大字号"""
        max_size = 0
        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                if size_pt > max_size:
                    max_size = size_pt
        return max_size

    def _is_isolated(self, doc, idx):
        """判断段落前后是否有空行"""
        prev_empty = idx > 0 and not doc.paragraphs[idx - 1].text.strip()
        next_empty = idx < len(doc.paragraphs) - 1 and not doc.paragraphs[idx + 1].text.strip()
        return prev_empty or next_empty

    def _is_isolated_standalone(self, para):
        """判断是否为独立短段落（近似判断）"""
        return len(para.text.strip()) <= 30

    def _first_nonempty_para(self, doc):
        """找到第一个非空段落索引"""
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                return i
        return 0

    def _set_paragraph_font(self, para, font_name, size_pt, bold=False):
        """设置段落所有 run 的字体和字号"""
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(size_pt)
            run.bold = bold
            # 设置东亚字体
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)

    def _set_line_spacing(self, para, spacing_pt):
        """设置固定行距"""
        pf = para.paragraph_format
        pf.line_spacing = Pt(spacing_pt)
        # 使用 XML 确保固定值
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._element.insert(0, pPr)
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")}></w:spacing>')
            pPr.append(spacing)
        spacing.set(qn('w:line'), str(int(spacing_pt * 20)))  # 转换为 twips
        spacing.set(qn('w:lineRule'), 'exact')

    def _set_spacing_before_after(self, para, before, after):
        """设置段前段后间距"""
        pf = para.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)

    def _set_first_line_indent(self, para, chars):
        """设置首行缩进（字符数）"""
        # 三号字16pt, 2字符 = 32pt = 640 twips
        indent_pt = chars * SIZE_SANHAO
        para.paragraph_format.first_line_indent = Pt(indent_pt)

    def _set_right_indent(self, para, chars):
        """设置右缩进（字符数）"""
        indent_pt = chars * SIZE_SANHAO
        para.paragraph_format.right_indent = Pt(indent_pt)

    def _set_left_indent(self, para, chars):
        """设置左缩进（字符数）"""
        indent_pt = chars * SIZE_SANHAO
        para.paragraph_format.left_indent = Pt(indent_pt)

    def _set_right_indent_chars(self, para, chars):
        """设置右缩进（字符数, 用于抄送等）"""
        indent_pt = chars * SIZE_SANHAO
        para.paragraph_format.right_indent = Pt(indent_pt)

    def _set_para_text(self, para, new_text):
        """安全替换段落文本，保留第一个 run 的格式"""
        if not para.runs:
            para.add_run(new_text)
            return
        # 保留第一个 run 的格式
        first_run = para.runs[0]
        # 清除所有 run
        for run in para.runs:
            run._element.getparent().remove(run._element)
        # 用第一个 run 的格式属性创建新 run
        new_run = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr>{self._run_rpr_xml(first_run)}</w:rPr>'
            f'  <w:t xml:space="preserve">{self._escape_xml(new_text)}</w:t>'
            f'</w:r>'
        )
        para._element.append(new_run)

    def _run_rpr_xml(self, run):
        """获取 run 的 rPr XML 字符串"""
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            return rPr.xml.replace(f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', '')
        return ''

    def _escape_xml(self, text):
        """XML 转义"""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    def _arabic_to_chinese(self, num):
        """阿拉伯数字转中文数字"""
        mapping = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
                   6: "六", 7: "七", 8: "八", 9: "九", 10: "十"}
        if num in mapping:
            return mapping[num]
        if num < 20:
            return "十" + (mapping.get(num - 10, "") if num > 10 else "")
        if num < 100:
            tens = num // 10
            ones = num % 10
            result = mapping.get(tens, "") + "十"
            if ones:
                result += mapping.get(ones, "")
            return result
        return str(num)

    def _chinese_to_arabic(self, cn):
        """中文数字转阿拉伯数字"""
        mapping = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                   "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
        if cn in mapping:
            return mapping[cn]
        return 1

    def _add_confidence(self, para_idx, element_type, confidence):
        """添加置信度标记"""
        emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}
        level = {"high": "高(90%+)", "medium": "中(60-90%)", "low": "低(<60%)"}
        self.confidence_marks.append({
            "para_index": para_idx,
            "element_type": element_type,
            "confidence": confidence,
            "emoji": emoji.get(confidence, "🔴"),
            "level": level.get(confidence, "低(<60%)")
        })
