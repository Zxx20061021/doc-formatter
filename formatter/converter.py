"""
文件格式转换模块
优先使用 LibreOffice headless，降级使用纯 Python 方案
"""

import os
import html
import subprocess
import logging

logger = logging.getLogger(__name__)

# LibreOffice 支持的转换格式映射
SUPPORTED_FORMATS = {
    "pdf": "pdf",
    "docx": "docx",
    "doc": "doc",
    "pptx": "pptx",
    "ppt": "ppt",
    "xlsx": "xlsx",
    "xls": "xls",
    "odt": "odt",
    "ods": "ods",
    "odp": "odp",
    "rtf": "rtf",
    "html": "html",
    "txt": "txt",
    "csv": "csv",
}

# LibreOffice 输出过滤器
OUTPUT_FILTERS = {
    "pdf": "writer_pdf_Export",
    "docx": "MS Word 2007 XML",
    "doc": "MS Word 97",
    "pptx": "Impress MS PowerPoint 2007 XML",
    "ppt": "MS PowerPoint 97",
    "xlsx": "Calc MS Excel 2007 XML",
    "xls": "MS Excel 97",
    "odt": "writer8",
    "ods": "calc8",
    "odp": "impress8",
    "rtf": "Rich Text Format",
    "html": "HTML (StarWriter)",
    "txt": "Text",
    "csv": "Text - txt - csv (StarCalc)",
}

# 纯 Python 降级方案支持的转换路径
FALLBACK_PATHS = {
    ('docx', 'html'): True,
    ('docx', 'txt'): True,
    ('docx', 'pdf'): True,   # 尝试 docx2pdf
    ('txt', 'html'): True,
    ('html', 'txt'): True,
}


def find_libreoffice():
    """查找 LibreOffice 可执行文件"""
    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/usr/local/bin/libreoffice",
        "/snap/bin/libreoffice",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    # 尝试 which
    try:
        result = subprocess.run(["which", "soffice"], capture_output=True, text=True, timeout=5)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
        result = subprocess.run(["which", "libreoffice"], capture_output=True, text=True, timeout=5)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass
    return None


def convert_file(input_path, output_format, output_dir=None):
    """
    转换文件格式。优先 LibreOffice，降级使用纯 Python 方案。

    Args:
        input_path: 输入文件路径
        output_format: 目标格式 (pdf, docx, pptx, xlsx 等)
        output_dir: 输出目录 (默认与输入同目录)

    Returns:
        dict: {"success": bool, "output_path": str, "error": str}
    """
    if output_format not in SUPPORTED_FORMATS:
        return {"success": False, "error": f"不支持的目标格式: {output_format}"}

    if output_dir is None:
        output_dir = os.path.dirname(input_path)

    # 尝试纯 Python 降级方案
    fallback = _try_fallback_convert(input_path, output_format, output_dir)
    if fallback is not None:
        return fallback

    # 使用 LibreOffice
    soffice = find_libreoffice()
    if not soffice:
        return {"success": False, "error": "此格式转换需要 LibreOffice 支持，服务器暂未安装。支持的纯 Python 转换：docx→html, docx→txt, txt→html, html→txt"}

    try:
        cmd = [
            soffice,
            "--headless",
            "--convert-to", output_format,
            "--outdir", output_dir,
            input_path
        ]

        logger.info(f"执行转换命令: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        if result.returncode != 0:
            logger.error(f"LibreOffice 转换失败: {result.stderr}")
            return {"success": False, "error": f"转换失败: {result.stderr}"}

        # 查找输出文件
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_dir, f"{base_name}.{output_format}")

        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path}
        else:
            for f in os.listdir(output_dir):
                if f.startswith(base_name) and f.endswith(f".{output_format}"):
                    return {"success": True, "output_path": os.path.join(output_dir, f)}

            return {"success": False, "error": "转换完成但找不到输出文件"}

    except subprocess.TimeoutExpired:
        return {"success": False, "error": "转换超时（120秒），文件可能过大"}
    except Exception as e:
        logger.exception("转换异常")
        return {"success": False, "error": f"转换异常: {str(e)}"}


def _try_fallback_convert(input_path, output_format, output_dir):
    """
    纯 Python 降级转换方案。
    支持: docx→html, docx→txt, docx→pdf(需docx2pdf), html→txt, txt→html
    """
    ext = os.path.splitext(input_path)[1].lstrip('.').lower()
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.{output_format}")

    # 检查是否有降级路径
    if (ext, output_format) not in FALLBACK_PATHS:
        return None

    try:
        if ext == 'docx' and output_format == 'html':
            return _docx_to_html(input_path, output_path)

        if ext == 'docx' and output_format == 'txt':
            return _docx_to_txt(input_path, output_path)

        if ext == 'docx' and output_format == 'pdf':
            return _docx_to_pdf(input_path, output_path)

        if ext == 'txt' and output_format == 'html':
            return _txt_to_html(input_path, output_path)

        if ext == 'html' and output_format == 'txt':
            return _html_to_txt(input_path, output_path)

        return None

    except Exception as e:
        logger.warning(f"降级转换失败: {e}")
        return None


def _docx_to_html(input_path, output_path):
    """docx 转 html（带 HTML 转义防止 XSS）"""
    from docx import Document
    doc = Document(input_path)

    html_parts = [
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<style>body{font-family:"SimSun","Noto Sans SC",sans-serif;max-width:800px;margin:0 auto;padding:20px;}',
        'p{margin:0.5em 0;line-height:1.8;} table{border-collapse:collapse;width:100%;}',
        'td,th{border:1px solid #ccc;padding:6px 8px;}</style></head><body>'
    ]

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            html_parts.append('<br>')
            continue

        style = ''
        align = para.alignment
        if align == 1:  # CENTER
            style += 'text-align:center;'
        elif align == 2:  # RIGHT
            style += 'text-align:right;'

        runs_html = ''
        for run in para.runs:
            rtext = html.escape(run.text)
            if run.bold:
                rtext = f'<b>{rtext}</b>'
            if run.italic:
                rtext = f'<i>{rtext}</i>'
            if run.underline:
                rtext = f'<u>{rtext}</u>'
            if run.font.size:
                size_pt = run.font.size.pt
                runs_html += f'<span style="font-size:{size_pt}pt;">{rtext}</span>'
            else:
                runs_html += rtext

        attr = f' style="{style}"' if style else ''
        html_parts.append(f'<p{attr}>{runs_html}</p>')

    # 表格
    for table in doc.tables:
        html_parts.append('<table>')
        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                html_parts.append(f'<td>{html.escape(cell.text)}</td>')
            html_parts.append('</tr>')
        html_parts.append('</table>')

    html_parts.append('</body></html>')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_parts))

    return {"success": True, "output_path": output_path}


def _docx_to_txt(input_path, output_path):
    """docx 转纯文本"""
    from docx import Document
    doc = Document(input_path)

    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')
        # 也提取表格文本
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                f.write('\t'.join(cells) + '\n')

    return {"success": True, "output_path": output_path}


def _docx_to_pdf(input_path, output_path):
    """docx 转 pdf，尝试 docx2pdf（需要安装）"""
    try:
        import docx2pdf
        docx2pdf.convert(input_path, output_path)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return {"success": True, "output_path": output_path}
    except ImportError:
        logger.info("docx2pdf 未安装，跳过 PDF 转换")
    except Exception as e:
        logger.warning(f"docx2pdf 转换失败: {e}")

    # 不使用 reportlab（中文字体问题太多），回退到 LibreOffice
    return None


def _txt_to_html(input_path, output_path):
    """txt 转 html"""
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()

    escaped = html.escape(text)
    html_content = (
        '<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<style>body{font-family:"Noto Sans SC",sans-serif;max-width:800px;margin:0 auto;padding:20px;'
        'white-space:pre-wrap;line-height:1.8;}</style></head><body>',
        f'{escaped}</body></html>'
    )

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(''.join(html_content))

    return {"success": True, "output_path": output_path}


def _html_to_txt(input_path, output_path):
    """html 转纯文本"""
    import re
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 移除 script 和 style 内容
    text = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    # 移除所有 HTML 标签
    text = re.sub(r'<[^>]+>', '', text)
    # 处理 HTML 实体
    text = html.unescape(text)
    # 压缩多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text.strip())

    return {"success": True, "output_path": output_path}


def get_supported_input_formats():
    """获取支持的输入格式列表"""
    return list(SUPPORTED_FORMATS.keys())


def get_conversion_options(source_ext):
    """根据源文件格式获取可转换的目标格式"""
    source_ext = source_ext.lstrip(".").lower()

    # 文档类
    doc_formats = ["docx", "doc", "odt", "rtf", "html", "txt", "pdf"]
    # 表格类
    sheet_formats = ["xlsx", "xls", "ods", "csv", "pdf"]
    # 演示类
    present_formats = ["pptx", "ppt", "odp", "pdf"]

    if source_ext in doc_formats:
        return [f for f in doc_formats if f != source_ext]
    elif source_ext in sheet_formats:
        return [f for f in sheet_formats if f != source_ext]
    elif source_ext in present_formats:
        return [f for f in present_formats if f != source_ext]
    else:
        return [f for f in SUPPORTED_FORMATS if f != source_ext]
