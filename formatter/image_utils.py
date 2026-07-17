"""
图片尺寸统一处理模块
用于处理 Word 文档中的图片，统一尺寸
"""

import os
import logging
from docx import Document
from docx.shared import Mm, Emu
from docx.oxml.ns import qn
from PIL import Image
import io

logger = logging.getLogger(__name__)

# A4 页面可用区域 (mm)
# A4: 210mm x 297mm, 边距: 上37 下35 左28 右26
A4_CONTENT_WIDTH_MM = 210 - 28 - 26   # 156mm
A4_CONTENT_HEIGHT_MM = 297 - 37 - 35   # 225mm

# 默认最大尺寸 (mm)
DEFAULT_MAX_WIDTH = A4_CONTENT_WIDTH_MM
DEFAULT_MAX_HEIGHT = A4_CONTENT_HEIGHT_MM


def mm_to_emu(mm_val):
    """毫米转 EMU (English Metric Units)"""
    return int(mm_val * 36000)


def emu_to_mm(emu_val):
    """EMU 转毫米"""
    if emu_val == 0:
        return 0
    return emu_val / 36000


def validate_dimensions(width_mm, height_mm):
    """
    验证图片尺寸是否在页面范围内

    Args:
        width_mm: 宽度（毫米）
        height_mm: 高度（毫米）

    Returns:
        dict: {"valid": bool, "error": str}
    """
    if width_mm <= 0 or height_mm <= 0:
        return {"valid": False, "error": "尺寸必须大于0"}

    if width_mm > A4_CONTENT_WIDTH_MM:
        return {"valid": False, "error": f"宽度 {width_mm}mm 超出页面可用宽度 {A4_CONTENT_WIDTH_MM}mm"}

    if height_mm > A4_CONTENT_HEIGHT_MM:
        return {"valid": False, "error": f"高度 {height_mm}mm 超出页面可用高度 {A4_CONTENT_HEIGHT_MM}mm"}

    return {"valid": True, "error": ""}


def resize_images_in_docx(input_path, output_path, target_width_mm=None, target_height_mm=None, mode="fit"):
    """
    统一 Word 文档中的图片尺寸

    Args:
        input_path: 输入 docx 文件路径
        output_path: 输出 docx 文件路径
        target_width_mm: 目标宽度（毫米），None 表示自动
        target_height_mm: 目标高度（毫米），None 表示自动
        mode: 处理模式
            - "fit": 按比例缩放到指定宽度内（保持宽高比）
            - "exact": 精确设置为指定尺寸
            - "max": 确保不超过页面范围

    Returns:
        dict: {"success": bool, "processed_count": int, "error": str}
    """
    try:
        doc = Document(input_path)
        processed = 0

        # 处理段落中的图片
        for para in doc.paragraphs:
            for run in para.runs:
                for drawing in run._element.findall(qn('w:drawing')):
                    processed += _process_drawing(drawing, target_width_mm, target_height_mm, mode)

        # 处理表格中的图片
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for drawing in run._element.findall(qn('w:drawing')):
                                processed += _process_drawing(drawing, target_width_mm, target_height_mm, mode)

        # 处理页眉页脚中的图片
        for section in doc.sections:
            for header_part in [section.header, section.first_page_header, section.even_page_header]:
                if header_part and header_part.is_linked_to_previous is False:
                    for para in header_part.paragraphs:
                        for run in para.runs:
                            for drawing in run._element.findall(qn('w:drawing')):
                                processed += _process_drawing(drawing, target_width_mm, target_height_mm, mode)
            for footer_part in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer_part and footer_part.is_linked_to_previous is False:
                    for para in footer_part.paragraphs:
                        for run in para.runs:
                            for drawing in run._element.findall(qn('w:drawing')):
                                processed += _process_drawing(drawing, target_width_mm, target_height_mm, mode)

        doc.save(output_path)
        return {"success": True, "processed_count": processed}

    except Exception as e:
        logger.exception("图片处理异常")
        return {"success": False, "error": str(e), "processed_count": 0}


def _process_drawing(drawing, target_width_mm, target_height_mm, mode):
    """处理单个 drawing 元素（支持 inline 和 anchor 类型）"""
    # 查找 extent 元素（inline 和 anchor 都有）
    extent = drawing.find('.//' + qn('wp:extent'))
    if extent is None:
        return 0

    # 获取当前尺寸 (EMU)
    current_cx = int(extent.get('cx', '0'))
    current_cy = int(extent.get('cy', '0'))

    if current_cx == 0 or current_cy == 0:
        return 0

    current_width_mm = emu_to_mm(current_cx)
    current_height_mm = emu_to_mm(current_cy)
    aspect = current_width_mm / current_height_mm

    new_width_mm = current_width_mm
    new_height_mm = current_height_mm

    if mode == "fit" and target_width_mm:
        # 按比例缩放到目标宽度
        new_width_mm = min(target_width_mm, DEFAULT_MAX_WIDTH)
        new_height_mm = new_width_mm / aspect
        if new_height_mm > DEFAULT_MAX_HEIGHT:
            new_height_mm = DEFAULT_MAX_HEIGHT
            new_width_mm = new_height_mm * aspect

    elif mode == "exact" and target_width_mm and target_height_mm:
        # 精确尺寸
        validation = validate_dimensions(target_width_mm, target_height_mm)
        if not validation["valid"]:
            return 0
        new_width_mm = target_width_mm
        new_height_mm = target_height_mm

    elif mode == "max":
        # 确保不超过页面范围
        if current_width_mm > DEFAULT_MAX_WIDTH:
            new_width_mm = DEFAULT_MAX_WIDTH
            new_height_mm = new_width_mm / aspect
        if new_height_mm > DEFAULT_MAX_HEIGHT:
            new_height_mm = DEFAULT_MAX_HEIGHT
            new_width_mm = new_height_mm * aspect

    # 更新 wp:extent 尺寸
    extent.set('cx', str(mm_to_emu(new_width_mm)))
    extent.set('cy', str(mm_to_emu(new_height_mm)))

    # 也更新 a:extent (如果存在)
    for a_extent in drawing.findall('.//' + qn('a:extent')):
        a_extent.set('cx', str(mm_to_emu(new_width_mm)))
        a_extent.set('cy', str(mm_to_emu(new_height_mm)))

    return 1


def get_document_images_info(input_path):
    """
    获取文档中所有图片的信息（包括段落、表格中的图片）

    Returns:
        list: [{"index": int, "width_mm": float, "height_mm": float, "location": str}]
    """
    try:
        doc = Document(input_path)
        images = []
        idx = 0

        # 段落中的图片
        for para in doc.paragraphs:
            for run in para.runs:
                for drawing in run._element.findall(qn('w:drawing')):
                    info = _get_drawing_info(drawing, idx, "段落")
                    if info:
                        images.append(info)
                        idx += 1

        # 表格中的图片
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for drawing in run._element.findall(qn('w:drawing')):
                                info = _get_drawing_info(drawing, idx, "表格")
                                if info:
                                    images.append(info)
                                    idx += 1

        return images
    except Exception as e:
        logger.exception("获取图片信息异常")
        return []


def _get_drawing_info(drawing, idx, location):
    """获取单个 drawing 的尺寸信息"""
    extent = drawing.find('.//' + qn('wp:extent'))
    if extent is not None:
        cx = int(extent.get('cx', '0'))
        cy = int(extent.get('cy', '0'))
        if cx > 0 and cy > 0:
            return {
                "index": idx,
                "width_mm": round(emu_to_mm(cx), 1),
                "height_mm": round(emu_to_mm(cy), 1),
                "location": location,
            }
    return None
